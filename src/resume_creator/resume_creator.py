import streamlit as st
from streamlit_extras.add_vertical_space import add_vertical_space
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import os
import time

class ResumeCreator:
    @staticmethod
    def create_resume_prompt(resume_data):
        return f'''Create a professional resume with the following information:

Personal Information:
Name: {resume_data['personal']['name']}
Email: {resume_data['personal']['email']}
Phone: {resume_data['personal']['phone']}
Location: {resume_data['personal']['location']}
LinkedIn: {resume_data['personal']['linkedin']}

Education:
{resume_data['education_text']}

Work Experience:
{resume_data['experience_text']}

Skills:
Technical Skills: {resume_data['skills']['technical']}
Soft Skills: {resume_data['skills']['soft']}

Projects:
{resume_data['projects_text']}

Certifications:
{resume_data['certifications_text']}

Format the resume in a modern, professional style with clear sections and bullet points.
Focus on achievements and quantifiable results.
Use action verbs and industry-specific keywords.
Ensure all dates are consistent and in reverse chronological order.
'''

    @staticmethod
    def suggest_improvements_prompt(section_type, content):
        return f'''As an AI resume expert, analyze and suggest improvements for this {section_type} section:

Current Content:
{content}

Provide specific suggestions to:
1. Make it more impactful and professional
2. Add quantifiable achievements
3. Use stronger action verbs
4. Include relevant keywords
5. Improve clarity and conciseness

Format your response as:
1. Improved Version: [provide enhanced text]
2. Key Improvements Made: [bullet points explaining changes]
'''

    @staticmethod
    def format_list_to_text(items):
        formatted_text = ""
        for item in items:
            if isinstance(item, dict):
                if 'institution' in item:  # Education
                    parts = [
                        f"• {item['institution']}" if item['institution'] else "",
                        f"{item['degree']}" if item['degree'] else "",
                        f"({item['year']})" if item['year'] else "",
                        f"GPA: {item['gpa']}" if item['gpa'] else "",
                        f"\n  {item['description']}" if item['description'] else ""
                    ]
                    formatted_text += " ".join(filter(None, parts)) + "\n\n"
                elif 'company' in item:  # Experience
                    parts = [
                        f"• {item['company']}" if item['company'] else "",
                        f"- {item['position']}" if item['position'] else "",
                        f"({item['duration']})" if item['duration'] else "",
                        f"\n  {item['description']}" if item['description'] else ""
                    ]
                    formatted_text += " ".join(filter(None, parts)) + "\n\n"
                elif 'name' in item and 'technologies' in item:  # Projects
                    parts = [
                        f"• {item['name']}" if item['name'] else "",
                        f"({item['technologies']})" if item['technologies'] else "",
                        f"\n  {item['description']}" if item['description'] else ""
                    ]
                    formatted_text += " ".join(filter(None, parts)) + "\n\n"
                elif 'name' in item and 'issuer' in item:  # Certifications
                    parts = [
                        f"• {item['name']}" if item['name'] else "",
                        f"from {item['issuer']}" if item['issuer'] else "",
                        f"({item['year']})" if item['year'] else "",
                        f"\n  {item['description']}" if item['description'] else ""
                    ]
                    formatted_text += " ".join(filter(None, parts)) + "\n\n"
        return formatted_text.strip()

    @staticmethod
    def add_section_header(doc, text, font_size=14, spacing_before=12, spacing_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(spacing_before)
        p.paragraph_format.space_after = Pt(spacing_after)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0x1E, 0x88, 0xE5)  # Professional blue color
        return p

    @staticmethod
    def add_horizontal_line(paragraph):
        p = paragraph._p
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1E88E5')
        pBdr.append(bottom)
        p.get_or_add_pPr().append(pBdr)

    @staticmethod
    def create_template():
        doc = Document()
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
        
        # Name section
        name = doc.add_paragraph()
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name.add_run('{{NAME}}')
        name_run.font.size = Pt(24)
        name_run.font.name = 'Calibri'
        name_run.font.color.rgb = RGBColor(0x00, 0x2D, 0x62)  # Dark blue
        
        # Contact Info
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact.add_run('{{EMAIL}} | {{PHONE}} | {{LOCATION}}\n{{LINKEDIN}}')
        contact_run.font.size = Pt(11)
        contact_run.font.name = 'Calibri'
        
        # Add horizontal line
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_break()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '002D62')
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        # Sections
        sections_data = [
            ('PROFESSIONAL SUMMARY', '{{SUMMARY}}'),
            ('EDUCATION', '{{EDUCATION}}'),
            ('PROFESSIONAL EXPERIENCE', '{{EXPERIENCE}}'),
            ('TECHNICAL SKILLS', '{{TECHNICAL_SKILLS}}'),
            ('SOFT SKILLS', '{{SOFT_SKILLS}}'),
            ('PROJECTS', '{{PROJECTS}}'),
            ('CERTIFICATIONS', '{{CERTIFICATIONS}}')
        ]
        
        for title, placeholder in sections_data:
            # Section Header
            header = doc.add_heading(level=1)
            header.alignment = WD_ALIGN_PARAGRAPH.LEFT
            header_run = header.add_run(title)
            header_run.font.name = 'Calibri'
            header_run.font.size = Pt(14)
            header_run.font.color.rgb = RGBColor(0x00, 0x2D, 0x62)
            header_run.font.bold = True
            
            # Add content placeholder
            content = doc.add_paragraph()
            content.add_run(placeholder)
            
            # Add space after section
            doc.add_paragraph()
        
        # Apply consistent formatting
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.paragraph_format.space_after = Pt(6)
            for run in paragraph.runs:
                if not run.font.size:
                    run.font.size = Pt(11)
                run.font.name = 'Calibri'
        
        return doc

    @staticmethod
    def generate_docx(resume_data):
        try:
            # Create document from template
            doc = ResumeCreator.create_template()
            
            # Replace placeholders with actual content
            for paragraph in doc.paragraphs:
                text = paragraph.text
                if '{{NAME}}' in text:
                    paragraph.text = text.replace('{{NAME}}', resume_data['personal']['name'])
                elif '{{EMAIL}}' in text:
                    contact_info = [
                        resume_data['personal']['email'],
                        resume_data['personal']['phone'],
                        resume_data['personal']['location']
                    ]
                    paragraph.text = text.replace('{{EMAIL}} | {{PHONE}} | {{LOCATION}}', ' | '.join(filter(None, contact_info)))
                elif '{{LINKEDIN}}' in text:
                    paragraph.text = text.replace('{{LINKEDIN}}', resume_data['personal']['linkedin'])
                elif '{{EDUCATION}}' in text:
                    paragraph.text = text.replace('{{EDUCATION}}', resume_data['education_text'])
                elif '{{EXPERIENCE}}' in text:
                    paragraph.text = text.replace('{{EXPERIENCE}}', resume_data['experience_text'])
                elif '{{TECHNICAL_SKILLS}}' in text:
                    paragraph.text = text.replace('{{TECHNICAL_SKILLS}}', resume_data['skills']['technical'])
                elif '{{SOFT_SKILLS}}' in text:
                    paragraph.text = text.replace('{{SOFT_SKILLS}}', resume_data['skills']['soft'])
                elif '{{PROJECTS}}' in text:
                    paragraph.text = text.replace('{{PROJECTS}}', resume_data['projects_text'])
                elif '{{CERTIFICATIONS}}' in text:
                    paragraph.text = text.replace('{{CERTIFICATIONS}}', resume_data['certifications_text'])

            # Save to BytesIO
            bio = io.BytesIO()
            doc.save(bio)
            bio.seek(0)
            return bio
        except Exception as e:
            st.error(f"Error generating DOCX: {str(e)}")
            return None

    @staticmethod
    def create_docx_resume(content):
        try:
            # Try to load template
            template_path = os.path.join(os.path.dirname(__file__), 'template.docx')
            if not os.path.exists(template_path):
                # Create and save template if it doesn't exist
                template = ResumeCreator.create_template()
                template.save(template_path)
            
            # Load template
            doc = Document(template_path)
            
            # Parse content sections
            sections = {}
            current_section = None
            current_content = []
            
            for line in content.split('\n'):
                if line.startswith('#'):
                    if current_section:
                        sections[current_section] = '\n'.join(current_content)
                    current_section = line.replace('#', '').strip()
                    current_content = []
                else:
                    current_content.append(line)
            
            if current_section:
                sections[current_section] = '\n'.join(current_content)
            
            # Replace placeholders in template
            for paragraph in doc.paragraphs:
                text = paragraph.text
                if '{{NAME}}' in text:
                    paragraph.text = text.replace('{{NAME}}', sections.get('Personal Information', '').split('\n')[0].strip())
                elif '{{EMAIL}}' in text:
                    info = sections.get('Personal Information', '').split('\n')
                    email = next((line for line in info if 'Email:' in line), '').replace('Email:', '').strip()
                    phone = next((line for line in info if 'Phone:' in line), '').replace('Phone:', '').strip()
                    location = next((line for line in info if 'Location:' in line), '').replace('Location:', '').strip()
                    linkedin = next((line for line in info if 'LinkedIn:' in line), '').replace('LinkedIn:', '').strip()
                    paragraph.text = text.replace('{{EMAIL}}', email).replace('{{PHONE}}', phone).replace('{{LOCATION}}', location)
                elif '{{LINKEDIN}}' in text:
                    paragraph.text = text.replace('{{LINKEDIN}}', linkedin)
                elif '{{EDUCATION}}' in text:
                    paragraph.text = text.replace('{{EDUCATION}}', sections.get('Education', ''))
                elif '{{EXPERIENCE}}' in text:
                    paragraph.text = text.replace('{{EXPERIENCE}}', sections.get('Work Experience', ''))
                elif '{{TECHNICAL_SKILLS}}' in text:
                    skills = sections.get('Skills', '').split('\n')
                    technical = next((line for line in skills if 'Technical Skills:' in line), '').replace('Technical Skills:', '').strip()
                    soft = next((line for line in skills if 'Soft Skills:' in line), '').replace('Soft Skills:', '').strip()
                    paragraph.text = text.replace('{{TECHNICAL_SKILLS}}', technical).replace('{{SOFT_SKILLS}}', soft)
                elif '{{PROJECTS}}' in text:
                    paragraph.text = text.replace('{{PROJECTS}}', sections.get('Projects', ''))
                elif '{{CERTIFICATIONS}}' in text:
                    paragraph.text = text.replace('{{CERTIFICATIONS}}', sections.get('Certifications', ''))
            
            return doc
        
        except Exception as e:
            st.error(f"Error using template: {str(e)}")
            # Fallback to original implementation
            return ResumeCreator._create_docx_resume_fallback(content)

    @staticmethod
    def _create_docx_resume_fallback(content):
        doc = Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
        
        sections = content.split('\n\n')
        for section in sections:
            if section.strip():
                if section.startswith('#'):
                    header = section.split('\n')[0].replace('#', '').strip()
                    p = ResumeCreator.add_section_header(doc, header)
                    ResumeCreator.add_horizontal_line(p)
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    p.paragraph_format.space_after = Pt(6)
                    for line in section.split('\n'):
                        if line.strip():
                            if line.strip().startswith('-'):
                                bullet_p = doc.add_paragraph(
                                    line.strip()[1:].strip(),
                                    style='List Bullet'
                                )
                                bullet_p.paragraph_format.left_indent = Inches(0.25)
                                bullet_p.paragraph_format.space_after = Pt(3)
                                bullet_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                                for run in bullet_p.runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(11)
                            else:
                                run = p.add_run(line.strip() + '\n')
                                run.font.name = 'Calibri'
                                run.font.size = Pt(11)
        return doc

    @staticmethod
    def get_ai_suggestions(gemini_api_key, section_type, content):
        try:
            generation_config = {
                "temperature": 0.7,
                "top_p": 1,
                "top_k": 1,
                "max_output_tokens": 1000,
            }
            model = genai.GenerativeModel(
                model_name="models/gemini-1.5-pro-latest",
                generation_config=generation_config
            )
            
            prompt = ResumeCreator.suggest_improvements_prompt(section_type, content)
            response = model.generate_content(prompt)
            return response.text
        except Exception as e:
            return f"Error getting suggestions: {str(e)}"

    @staticmethod
    def enhance_text(gemini_api_key, text, section_type):
        try:
            time.sleep(1)
            
            generation_config = {
                "temperature": 0.7,
                "top_p": 1,
                "top_k": 1,
                "max_output_tokens": 1000,
            }

            genai.configure(api_key=gemini_api_key)
            
            model = genai.GenerativeModel(
                model_name="models/gemini-1.5-pro-latest",
                generation_config=generation_config
            )
            
            prompt = f'''Enhance this {section_type} description to be more professional and impactful:

{text}

Make it:
1. More impactful with quantifiable achievements
2. Use strong action verbs
3. Include relevant keywords
4. Clear and concise
5. Professional tone

Return ONLY the enhanced text without any explanations.'''
            
            response = model.generate_content(prompt)
            return response.text.strip()
        except Exception as e:
            if "429" in str(e):
                st.error("API rate limit reached. Please wait a few minutes and try again.")
            else:
                st.error(f"Error enhancing text: {str(e)}")
            return text

    @staticmethod
    def generate_resume_content(gemini_api_key, resume_data):
        try:
            time.sleep(1)
            
            genai.configure(api_key=gemini_api_key)
            generation_config = {
                "temperature": 0.7,
                "top_p": 1,
                "top_k": 1,
                "max_output_tokens": 2048,
            }
            model = genai.GenerativeModel(
                model_name="models/gemini-1.5-pro-latest",
                generation_config=generation_config
            )
            
            prompt = ResumeCreator.create_resume_prompt(resume_data)
            response = model.generate_content(prompt)
            return response.text.strip()
        except Exception as e:
            if "429" in str(e):
                st.error("API rate limit reached. Please wait a few minutes and try again.")
            else:
                st.error(f"Error generating resume: {str(e)}")
            return None

    @staticmethod
    def initialize_session_state():
        if 'resume_data' not in st.session_state:
            st.session_state.resume_data = {
                'personal': {
                    'name': '',
                    'email': '',
                    'phone': '',
                    'location': '',
                    'linkedin': ''
                },
                'education': [{'institution': '', 'degree': '', 'year': '', 'gpa': '', 'description': ''}],
                'experience': [{'company': '', 'position': '', 'duration': '', 'description': ''}],
                'skills': {
                    'technical': '',
                    'soft': ''
                },
                'projects': [{'name': '', 'technologies': '', 'description': ''}],
                'certifications': [{'name': '', 'issuer': '', 'year': '', 'description': ''}]
            }
        if 'resume_preview' not in st.session_state:
            st.session_state.resume_preview = None
        if 'edit_mode' not in st.session_state:
            st.session_state.edit_mode = False

    @staticmethod
    def add_entry(section_name):
        if section_name == 'education':
            st.session_state.resume_data['education'].append(
                {'institution': '', 'degree': '', 'year': '', 'gpa': '', 'description': ''})
        elif section_name == 'experience':
            st.session_state.resume_data['experience'].append(
                {'company': '', 'position': '', 'duration': '', 'description': ''})
        elif section_name == 'projects':
            st.session_state.resume_data['projects'].append(
                {'name': '', 'technologies': '', 'description': ''})
        elif section_name == 'certifications':
            st.session_state.resume_data['certifications'].append(
                {'name': '', 'issuer': '', 'year': '', 'description': ''})
        st.rerun()

    @staticmethod
    def create_resume(gemini_api_key):
        st.markdown(
            """
            <style>
            .stTextArea textarea {
                font-size: 14px !important;
            }
            </style>
            """,
            unsafe_allow_html=True
        )
        
        st.markdown(f'<h4 style="text-align: center;">Create Your Professional Resume</h4>', 
                  unsafe_allow_html=True)
        
        ResumeCreator.initialize_session_state()
        
        # Personal Information Section
        st.markdown("### 📋 Personal Information")
        with st.form(key='personal_form'):
            col1, col2 = st.columns(2)
            with col1:
                st.session_state.resume_data['personal']['name'] = st.text_input(
                    "Full Name", value=st.session_state.resume_data['personal']['name'])
                st.session_state.resume_data['personal']['email'] = st.text_input(
                    "Email", value=st.session_state.resume_data['personal']['email'])
                st.session_state.resume_data['personal']['phone'] = st.text_input(
                    "Phone", value=st.session_state.resume_data['personal']['phone'])
            with col2:
                st.session_state.resume_data['personal']['location'] = st.text_input(
                    "Location", value=st.session_state.resume_data['personal']['location'])
                st.session_state.resume_data['personal']['linkedin'] = st.text_input(
                    "LinkedIn Profile", value=st.session_state.resume_data['personal']['linkedin'])
            st.form_submit_button("Save Personal Info")

        # Education Section
        st.markdown("### 🎓 Education")
        for i, edu in enumerate(st.session_state.resume_data['education']):
            with st.expander(f"Education Entry {i+1}", expanded=True):
                with st.form(key=f'education_form_{i}'):
                    col1, col2 = st.columns(2)
                    with col1:
                        edu['institution'] = st.text_input(
                            "Institution", key=f"edu_inst_{i}", 
                            value=edu['institution'])
                        edu['degree'] = st.text_input(
                            "Degree", key=f"edu_deg_{i}",
                            value=edu['degree'])
                    with col2:
                        edu['year'] = st.text_input(
                            "Year", key=f"edu_year_{i}",
                            value=edu['year'])
                        edu['gpa'] = st.text_input(
                            "GPA", key=f"edu_gpa_{i}",
                            value=edu['gpa'])
                    
                    desc_key = f"edu_desc_{i}"
                    edu['description'] = st.text_area(
                        "Additional Details", key=desc_key,
                        value=edu['description'])
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.form_submit_button("Save"):
                            st.success("Education details saved!")
                    with col2:
                        if st.form_submit_button("✨ Enhance with AI"):
                            enhanced_text = ResumeCreator.enhance_text(
                                gemini_api_key, edu['description'], "Education"
                            )
                            edu['description'] = enhanced_text
                            st.rerun()

        # Work Experience Section
        st.markdown("### 💼 Work Experience")
        for i, exp in enumerate(st.session_state.resume_data['experience']):
            with st.expander(f"Experience Entry {i+1}", expanded=True):
                with st.form(key=f'experience_form_{i}'):
                    col1, col2 = st.columns(2)
                    with col1:
                        exp['company'] = st.text_input(
                            "Company", key=f"exp_comp_{i}",
                            value=exp['company'])
                        exp['position'] = st.text_input(
                            "Position", key=f"exp_pos_{i}",
                            value=exp['position'])
                    with col2:
                        exp['duration'] = st.text_input(
                            "Duration", key=f"exp_dur_{i}",
                            value=exp['duration'])
                    
                    desc_key = f"exp_desc_{i}"
                    exp['description'] = st.text_area(
                        "Responsibilities and Achievements", key=desc_key,
                        value=exp['description'])
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.form_submit_button("Save"):
                            st.success("Experience details saved!")
                    with col2:
                        if st.form_submit_button("✨ Enhance with AI"):
                            enhanced_text = ResumeCreator.enhance_text(
                                gemini_api_key, exp['description'], "Work Experience"
                            )
                            exp['description'] = enhanced_text
                            st.rerun()

        # Skills Section
        st.markdown("### 🔧 Skills")
        with st.form(key='skills_form'):
            col1, col2 = st.columns(2)
            with col1:
                tech_key = "technical_skills"
                st.session_state.resume_data['skills']['technical'] = st.text_area(
                    "Technical Skills", value=st.session_state.resume_data['skills']['technical'],
                    key=tech_key)
                
                if st.form_submit_button("✨ Enhance Technical Skills"):
                    enhanced_text = ResumeCreator.enhance_text(
                        gemini_api_key, 
                        st.session_state.resume_data['skills']['technical'],
                        "Technical Skills"
                    )
                    st.session_state.resume_data['skills']['technical'] = enhanced_text
                    st.rerun()
            
            with col2:
                soft_key = "soft_skills"
                st.session_state.resume_data['skills']['soft'] = st.text_area(
                    "Soft Skills", value=st.session_state.resume_data['skills']['soft'],
                    key=soft_key)
                
                if st.form_submit_button("✨ Enhance Soft Skills"):
                    enhanced_text = ResumeCreator.enhance_text(
                        gemini_api_key, 
                        st.session_state.resume_data['skills']['soft'],
                        "Soft Skills"
                    )
                    st.session_state.resume_data['skills']['soft'] = enhanced_text
                    st.rerun()
            
            if st.form_submit_button("Save Skills"):
                st.success("Skills saved!")

        # Projects Section
        st.markdown("### 🚀 Projects")
        for i, proj in enumerate(st.session_state.resume_data['projects']):
            with st.expander(f"Project {i+1}", expanded=True):
                with st.form(key=f'project_form_{i}'):
                    col1, col2 = st.columns(2)
                    with col1:
                        proj['name'] = st.text_input(
                            "Project Name", key=f"proj_name_{i}",
                            value=proj['name'])
                    with col2:
                        proj['technologies'] = st.text_input(
                            "Technologies Used", key=f"proj_tech_{i}",
                            value=proj['technologies'])
                    
                    desc_key = f"proj_desc_{i}"
                    proj['description'] = st.text_area(
                        "Project Description", key=desc_key,
                        value=proj['description'])
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.form_submit_button("Save"):
                            st.success("Project details saved!")
                    with col2:
                        if st.form_submit_button("✨ Enhance with AI"):
                            enhanced_text = ResumeCreator.enhance_text(
                                gemini_api_key, proj['description'], "Project"
                            )
                            proj['description'] = enhanced_text
                            st.rerun()

        # Certifications Section
        st.markdown("### 🏆 Certifications")
        for i, cert in enumerate(st.session_state.resume_data['certifications']):
            with st.expander(f"Certification {i+1}", expanded=True):
                with st.form(key=f'certification_form_{i}'):
                    col1, col2 = st.columns(2)
                    with col1:
                        cert['name'] = st.text_input(
                            "Certification Name", key=f"cert_name_{i}",
                            value=cert['name'])
                        cert['issuer'] = st.text_input(
                            "Issuing Organization", key=f"cert_org_{i}",
                            value=cert['issuer'])
                    with col2:
                        cert['year'] = st.text_input(
                            "Year", key=f"cert_year_{i}",
                            value=cert['year'])
                    
                    desc_key = f"cert_desc_{i}"
                    cert['description'] = st.text_area(
                        "Description", key=desc_key,
                        value=cert['description'])
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.form_submit_button("Save"):
                            st.success("Certification details saved!")
                    with col2:
                        if st.form_submit_button("✨ Enhance with AI"):
                            enhanced_text = ResumeCreator.enhance_text(
                                gemini_api_key, cert['description'], "Certification"
                            )
                            cert['description'] = enhanced_text
                            st.rerun()

        # Add Entry Buttons
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            if st.button("➕ Add Education"):
                ResumeCreator.add_entry('education')
                st.rerun()
        with col2:
            if st.button("➕ Add Experience"):
                ResumeCreator.add_entry('experience')
                st.rerun()
        with col3:
            if st.button("➕ Add Project"):
                ResumeCreator.add_entry('projects')
                st.rerun()
        with col4:
            if st.button("➕ Add Certification"):
                ResumeCreator.add_entry('certifications')
                st.rerun()

        # Final Generate Button
        with st.form(key='generate_form'):
            generate_clicked = st.form_submit_button("Generate Resume", type="primary")
            
        if generate_clicked:
            try:
                with st.spinner('Creating your professional resume...'):
                    # Format text for sections with multiple entries
                    education_text = ResumeCreator.format_list_to_text(st.session_state.resume_data['education'])
                    experience_text = ResumeCreator.format_list_to_text(st.session_state.resume_data['experience'])
                    projects_text = ResumeCreator.format_list_to_text(st.session_state.resume_data['projects'])
                    certifications_text = ResumeCreator.format_list_to_text(st.session_state.resume_data['certifications'])
                    
                    # Add formatted text to resume data
                    resume_data = st.session_state.resume_data.copy()
                    resume_data['education_text'] = education_text
                    resume_data['experience_text'] = experience_text
                    resume_data['projects_text'] = projects_text
                    resume_data['certifications_text'] = certifications_text

                    # Generate DOCX file
                    bio = ResumeCreator.generate_docx(resume_data)
                    if bio:
                        # Create download button
                        file_name = f"{resume_data['personal']['name'].replace(' ', '_')}_Resume.docx"
                        st.download_button(
                            label="📥 Download Resume",
                            data=bio.getvalue(),
                            file_name=file_name,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        st.success("Resume generated successfully! Click the button above to download.")
            
            except Exception as e:
                st.error(f"Error generating resume: {str(e)}")
                return

    @staticmethod
    def show_preview_and_edit_options():
        if st.session_state.resume_preview:
            st.markdown("### Resume Preview")
            
            if st.session_state.edit_mode:
                edited_content = st.text_area("Edit your resume content",
                                            value=st.session_state.resume_preview,
                                            height=500)
                st.session_state.resume_preview = edited_content
            
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("Edit Resume"):
                    st.session_state.edit_mode = True
            
            with col2:
                if st.button("Download Resume"):
                    try:
                        doc = ResumeCreator.create_docx_resume(st.session_state.resume_preview)
                        docx_file = io.BytesIO()
                        doc.save(docx_file)
                        docx_file.seek(0)
                        
                        st.download_button(
                            label="Click to Download",
                            data=docx_file,
                            file_name=f"{st.session_state.resume_data['personal']['name'].lower().replace(' ', '_')}_resume.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    
                    except Exception as e:
                        st.error(f"Error creating DOCX: {str(e)}")
            
            with st.expander("💡 Resume Tips"):
                st.markdown("""
                - Use action verbs to start bullet points
                - Quantify achievements when possible
                - Keep formatting consistent
                - Proofread for spelling and grammar
                - Customize for each job application
                - Include relevant keywords from job descriptions
                - Keep it concise (1-2 pages)
                """)

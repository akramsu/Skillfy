import time
import numpy as np
import pandas as pd
import streamlit as st
import os
from dotenv import load_dotenv
from streamlit_option_menu import option_menu
from streamlit_extras.add_vertical_space import add_vertical_space
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain.vectorstores import FAISS
import google.generativeai as genai
from langchain.chains.question_answering import load_qa_chain
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.common.exceptions import NoSuchElementException
import warnings
warnings.filterwarnings('ignore')

# Load environment variables
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

# Configure Gemini API globally with explicit credentials
os.environ["GOOGLE_API_KEY"] = GEMINI_API_KEY
genai.configure(api_key=GEMINI_API_KEY)

def streamlit_config():

    # page configuration
    st.set_page_config(page_title='Resume Analyzer AI', layout="wide")

    # page header transparent color
    page_background_color = """
    <style>

    [data-testid="stHeader"] 
    {
    background: rgba(0,0,0,0);
    }

    </style>
    """
    st.markdown(page_background_color, unsafe_allow_html=True)

    # title and position
    st.markdown(f'<h1 style="text-align: center;">Resume Analyzer AI</h1>',
                unsafe_allow_html=True)


class resume_analyzer:

    def pdf_to_chunks(pdf):
        # read pdf and it returns memory address
        pdf_reader = PdfReader(pdf)

        # extrat text from each page separately
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()

        # Split the long text into small chunks.
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=700,
            chunk_overlap=200,
            length_function=len)

        chunks = text_splitter.split_text(text=text)
        return chunks


    def gemini(gemini_api_key, chunks, analyze):
        try:
            # Using Gemini service for embedding with explicit credentials
            os.environ["GOOGLE_API_KEY"] = gemini_api_key  # Set environment variable for this session
            
            # Initialize embeddings with the environment variable approach
            embeddings = GoogleGenerativeAIEmbeddings(
                model="models/embedding-001",
            )

            # Facebook AI Similarity Search library help us to convert text data to numerical vector
            vectorstores = FAISS.from_texts(chunks, embedding=embeddings)

            # compares the query and chunks, enabling the selection of the top 'K' most similar chunks based on their similarity scores.
            docs = vectorstores.similarity_search(query=analyze, k=3)
            
            # Extract the content from the documents
            doc_content = "\n\n".join([doc.page_content for doc in docs])
            
            # Set up the model
            generation_config = {
                "temperature": 0.7,
                "top_p": 1,
                "top_k": 1,
                "max_output_tokens": 1000,
            }
            
            # Create the model without explicitly passing the API key (it will use the environment variable)
            model = genai.GenerativeModel(
                model_name="models/gemini-1.5-pro-latest",
                generation_config=generation_config,
            )
            
            # Combine the query with the document content
            prompt = f"""
            Based on the following resume information:
            
            {doc_content}
            
            Please answer the following question:
            {analyze}
            """
            
            # Generate a response using Gemini
            response = model.generate_content(prompt)
            
            # Return the generated text
            return response.text
            
        except Exception as e:
            # Print detailed error information
            print(f"Error with Gemini API: {str(e)}")
            # Return a fallback message
            return f"Error with Gemini API: {str(e)}"


    def summary_prompt(query_with_chunks):
        query = f'''Provide a comprehensive and detailed analysis of the following resume. 
        
        Structure your response in the following format:
        1. **Professional Summary**: A concise overview of the candidate's background, experience, and key qualifications.
        2. **Key Skills**: List the top technical and soft skills evident in the resume.
        3. **Experience Highlights**: Summarize the most significant professional experiences and achievements.
        4. **Education & Certifications**: Highlight educational background and relevant certifications.
        5. **Overall Assessment**: Provide a final evaluation of the candidate's profile and potential fit for roles.

        Format each section with proper headings and bullet points for better readability.

        Resume Content:
        """""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""
        {query_with_chunks}
        """""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""
        '''
        return query


    def resume_summary():

        with st.form(key='Summary'):

            # User Upload the Resume
            add_vertical_space(1)
            pdf = st.file_uploader(label='Upload Your Resume', type='pdf')
            add_vertical_space(1)

            # Hidden Gemini API Key
            gemini_api_key = GEMINI_API_KEY
            add_vertical_space(2)

            # Click on Submit Button
            submit = st.form_submit_button(label='Submit')
            add_vertical_space(1)
        
        add_vertical_space(3)
        if submit:
            if pdf is not None and gemini_api_key != '':
                try:
                    with st.spinner('Processing...'):

                        pdf_chunks = resume_analyzer.pdf_to_chunks(pdf)

                        summary_prompt = resume_analyzer.summary_prompt(query_with_chunks=pdf_chunks)

                        summary = resume_analyzer.gemini(gemini_api_key=gemini_api_key, chunks=pdf_chunks, analyze=summary_prompt)

                    st.markdown(f'<h3 style="color: #1E88E5; border-bottom: 2px solid #1E88E5; padding-bottom: 8px;">Resume Analysis</h3>', unsafe_allow_html=True)
                    st.markdown(summary, unsafe_allow_html=True)

                except Exception as e:
                    st.markdown(f'<h5 style="text-align: center;color: orange;">{e}</h5>', unsafe_allow_html=True)

            elif pdf is None:
                st.markdown(f'<h5 style="text-align: center;color: orange;">Please Upload Your Resume</h5>', unsafe_allow_html=True)
            
            elif gemini_api_key == '':
                st.markdown(f'<h5 style="text-align: center;color: orange;">Gemini API Key not found in environment variables</h5>', unsafe_allow_html=True)


    def strength_prompt(query_with_chunks):
        query = f'''Analyze the strengths in the following resume and provide a detailed assessment.
        
        Structure your response in the following format:
        1. **Technical Strengths**: Identify and analyze the technical skills, tools, and technologies the candidate excels in.
        2. **Professional Strengths**: Highlight strong professional experiences, achievements, and responsibilities.
        3. **Soft Skills & Qualities**: Identify communication, leadership, problem-solving, and other soft skills evident in the resume.
        4. **Educational Strengths**: Note any impressive educational qualifications, certifications, or specialized training.
        5. **Unique Selling Points**: Identify what makes this candidate stand out from others in their field.
        
        Format each section with proper headings and bullet points for better readability.
        Be constructive and supportive in your feedback.
        
        Resume Content:
        """""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""
        {query_with_chunks}
        """""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""
        '''
        return query


    def resume_strength():

        with st.form(key='Strength'):

            # User Upload the Resume
            add_vertical_space(1)
            pdf = st.file_uploader(label='Upload Your Resume', type='pdf')
            add_vertical_space(1)

            # Hidden Gemini API Key
            gemini_api_key = GEMINI_API_KEY
            add_vertical_space(2)

            # Click on Submit Button
            submit = st.form_submit_button(label='Submit')
            add_vertical_space(1)
        
        add_vertical_space(3)
        if submit:
            if pdf is not None and gemini_api_key != '':
                try:
                    with st.spinner('Processing...'):

                        pdf_chunks = resume_analyzer.pdf_to_chunks(pdf)

                        strength_prompt = resume_analyzer.strength_prompt(query_with_chunks=pdf_chunks)

                        strength = resume_analyzer.gemini(gemini_api_key=gemini_api_key, chunks=pdf_chunks, analyze=strength_prompt)

                    st.markdown(f'<h3 style="color: #4CAF50; border-bottom: 2px solid #4CAF50; padding-bottom: 8px;">Resume Strengths Analysis</h3>', unsafe_allow_html=True)
                    st.markdown(strength, unsafe_allow_html=True)

                except Exception as e:
                    st.markdown(f'<h5 style="text-align: center;color: orange;">{e}</h5>', unsafe_allow_html=True)

            elif pdf is None:
                st.markdown(f'<h5 style="text-align: center;color: orange;">Please Upload Your Resume</h5>', unsafe_allow_html=True)
            
            elif gemini_api_key == '':
                st.markdown(f'<h5 style="text-align: center;color: orange;">Gemini API Key not found in environment variables</h5>', unsafe_allow_html=True)


    def weakness_prompt(query_with_chunks):
        query = f'''Analyze the potential areas for improvement in the following resume and provide constructive feedback.
        
        Structure your response in the following format:
        1. **Content Gaps**: Identify missing information or sections that would strengthen the resume content.
        2. **Skills Development**: Suggest skills that could be developed or emphasized based on industry trends.
        3. **Experience Presentation**: Highlight how experiences could be better presented or quantified.
        4. **Format & Structure**: Provide feedback on the organization, clarity, and presentation of information.
        5. **Improvement Recommendations**: Offer specific, actionable suggestions to address each weakness.
        
        Format each section with proper headings and bullet points for better readability.
        Be constructive and supportive in your feedback.
        
        Resume Content:
        """""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""
        {query_with_chunks}
        """""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""
        '''
        return query


    def resume_weakness():

        with st.form(key='Weakness'):

            # User Upload the Resume
            add_vertical_space(1)
            pdf = st.file_uploader(label='Upload Your Resume', type='pdf')
            add_vertical_space(1)

            # Hidden Gemini API Key
            gemini_api_key = GEMINI_API_KEY
            add_vertical_space(2)

            # Click on Submit Button
            submit = st.form_submit_button(label='Submit')
            add_vertical_space(1)
        
        add_vertical_space(3)
        if submit:
            if pdf is not None and gemini_api_key != '':
                try:
                    with st.spinner('Processing...'):

                        pdf_chunks = resume_analyzer.pdf_to_chunks(pdf)

                        weakness_prompt = resume_analyzer.weakness_prompt(query_with_chunks=pdf_chunks)

                        weakness = resume_analyzer.gemini(gemini_api_key=gemini_api_key, chunks=pdf_chunks, analyze=weakness_prompt)

                    st.markdown(f'<h3 style="color: #FF5722; border-bottom: 2px solid #FF5722; padding-bottom: 8px;">Areas for Improvement</h3>', unsafe_allow_html=True)
                    st.markdown(weakness, unsafe_allow_html=True)

                except Exception as e:
                    st.markdown(f'<h5 style="text-align: center;color: orange;">{e}</h5>', unsafe_allow_html=True)

            elif pdf is None:
                st.markdown(f'<h5 style="text-align: center;color: orange;">Please Upload Your Resume</h5>', unsafe_allow_html=True)
            
            elif gemini_api_key == '':
                st.markdown(f'<h5 style="text-align: center;color: orange;">Gemini API Key not found in environment variables</h5>', unsafe_allow_html=True)


    def improvement_prompt(query_with_chunks):
        query = f'''Provide detailed recommendations to improve the following resume and make it more effective.
        
        Structure your response in the following format:
        1. **Content Enhancements**: Suggest specific additions or modifications to strengthen the resume content.
        2. **Skills Presentation**: Recommend how to better showcase skills and align them with industry expectations.
        3. **Achievement Highlighting**: Provide guidance on how to better quantify and emphasize achievements.
        4. **Format & Structure Improvements**: Suggest layout, organization, and visual presentation improvements.
        5. **ATS Optimization**: Recommend changes to make the resume more ATS-friendly.
        6. **Industry-Specific Recommendations**: Provide tailored advice based on the candidate's industry or target roles.
        
        Format each section with proper headings and bullet points for better readability.
        Provide specific, actionable advice that can be immediately implemented.
        
        Resume Content:
        """""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""
        {query_with_chunks}
        """""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""
        '''
        return query


    def resume_improvement():

        with st.form(key='Improvement'):

            # User Upload the Resume
            add_vertical_space(1)
            pdf = st.file_uploader(label='Upload Your Resume', type='pdf')
            add_vertical_space(1)

            # Hidden Gemini API Key
            gemini_api_key = GEMINI_API_KEY
            add_vertical_space(2)

            # Click on Submit Button
            submit = st.form_submit_button(label='Submit')
            add_vertical_space(1)
        
        add_vertical_space(3)
        if submit:
            if pdf is not None and gemini_api_key != '':
                try:
                    with st.spinner('Processing...'):

                        pdf_chunks = resume_analyzer.pdf_to_chunks(pdf)

                        improvement_prompt = resume_analyzer.improvement_prompt(query_with_chunks=pdf_chunks)

                        improvement = resume_analyzer.gemini(gemini_api_key=gemini_api_key, chunks=pdf_chunks, analyze=improvement_prompt)

                    st.markdown(f'<h3 style="color: #9C27B0; border-bottom: 2px solid #9C27B0; padding-bottom: 8px;">Resume Improvement Recommendations</h3>', unsafe_allow_html=True)
                    st.markdown(improvement, unsafe_allow_html=True)

                except Exception as e:
                    st.markdown(f'<h5 style="text-align: center;color: orange;">{e}</h5>', unsafe_allow_html=True)

            elif pdf is None:
                st.markdown(f'<h5 style="text-align: center;color: orange;">Please Upload Your Resume</h5>', unsafe_allow_html=True)
            
            elif gemini_api_key == '':
                st.markdown(f'<h5 style="text-align: center;color: orange;">Gemini API Key not found in environment variables</h5>', unsafe_allow_html=True)


    def ats_prompt(query_with_chunks, job_description):
        query = f'''Analyze how well the following resume matches the provided job description and provide an ATS (Applicant Tracking System) compatibility score.
        
        Structure your response in the following format:
        1. **Overall ATS Score**: Provide a score out of 100 and explain the rationale.
        2. **Keyword Analysis**: Identify which important keywords from the job description are present or missing in the resume.
        3. **Skills Match**: Analyze how well the candidate's skills align with the job requirements.
        4. **Experience Relevance**: Evaluate how relevant the candidate's experience is to the job description.
        5. **Improvement Recommendations**: Provide specific suggestions to increase the ATS score and better match the job description.
        
        Format each section with proper headings and bullet points for better readability.
        
        Resume Content:
        """""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""
        {query_with_chunks}
        """""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""
        
        Job Description:
        """""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""
        {job_description}
        """""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""
        '''
        return query


    def resume_ats():

        with st.form(key='ATS'):

            # User Upload the Resume
            add_vertical_space(1)
            pdf = st.file_uploader(label='Upload Your Resume', type='pdf')
            add_vertical_space(1)

            # User Enter Job Description
            job_description = st.text_area(label='Enter Job Description')
            add_vertical_space(1)

            # Hidden Gemini API Key
            gemini_api_key = GEMINI_API_KEY
            add_vertical_space(2)

            # Click on Submit Button
            submit = st.form_submit_button(label='Submit')
            add_vertical_space(1)
        
        add_vertical_space(3)
        if submit:
            if pdf is not None and job_description != '' and gemini_api_key != '':
                try:
                    with st.spinner('Processing...'):

                        pdf_chunks = resume_analyzer.pdf_to_chunks(pdf)

                        ats_prompt = resume_analyzer.ats_prompt(query_with_chunks=pdf_chunks, job_description=job_description)

                        ats = resume_analyzer.gemini(gemini_api_key=gemini_api_key, chunks=pdf_chunks, analyze=ats_prompt)

                    st.markdown(f'<h3 style="color: #2196F3; border-bottom: 2px solid #2196F3; padding-bottom: 8px;">ATS Compatibility Analysis</h3>', unsafe_allow_html=True)
                    st.markdown(ats, unsafe_allow_html=True)

                except Exception as e:
                    st.markdown(f'<h5 style="text-align: center;color: orange;">{e}</h5>', unsafe_allow_html=True)

            elif pdf is None:
                st.markdown(f'<h5 style="text-align: center;color: orange;">Please Upload Your Resume</h5>', unsafe_allow_html=True)
            
            elif job_description == '':
                st.markdown(f'<h5 style="text-align: center;color: orange;">Please Enter Job Description</h5>', unsafe_allow_html=True)
            
            elif gemini_api_key == '':
                st.markdown(f'<h5 style="text-align: center;color: orange;">Gemini API Key not found in environment variables</h5>', unsafe_allow_html=True)


    def job_title_prompt(query_with_chunks):
        query = f'''Based on the following resume, suggest suitable job titles and roles that match the candidate's skills and experience.
        
        Structure your response in the following format:
        1. **Primary Job Titles**: List 3-5 job titles that are the best match for this candidate's profile.
        2. **Industry Recommendations**: Suggest 2-3 industries where this candidate would be most competitive.
        3. **Career Level**: Indicate whether the resume suggests entry-level, mid-level, senior, or executive positions.
        4. **Specialization Areas**: Identify niche areas or specializations where the candidate could excel.
        5. **Growth Opportunities**: Suggest potential career progression paths based on the current profile.
        
        For each job title, provide a brief explanation of why it's a good match.
        Format each section with proper headings and bullet points for better readability.
        
        Resume Content:
        """""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""
        {query_with_chunks}
        """""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""
        '''
        return query


    def resume_job_title():

        with st.form(key='Job Title'):

            # User Upload the Resume
            add_vertical_space(1)
            pdf = st.file_uploader(label='Upload Your Resume', type='pdf')
            add_vertical_space(1)

            # Hidden Gemini API Key
            gemini_api_key = GEMINI_API_KEY
            add_vertical_space(2)

            # Click on Submit Button
            submit = st.form_submit_button(label='Submit')
            add_vertical_space(1)
        
        add_vertical_space(3)
        if submit:
            if pdf is not None and gemini_api_key != '':
                try:
                    with st.spinner('Processing...'):

                        pdf_chunks = resume_analyzer.pdf_to_chunks(pdf)

                        job_title_prompt = resume_analyzer.job_title_prompt(query_with_chunks=pdf_chunks)

                        job_title = resume_analyzer.gemini(gemini_api_key=gemini_api_key, chunks=pdf_chunks, analyze=job_title_prompt)

                    st.markdown(f'<h3 style="color: #FFC107; border-bottom: 2px solid #FFC107; padding-bottom: 8px;">Recommended Job Titles & Career Paths</h3>', unsafe_allow_html=True)
                    st.markdown(job_title, unsafe_allow_html=True)

                except Exception as e:
                    st.markdown(f'<h5 style="text-align: center;color: orange;">{e}</h5>', unsafe_allow_html=True)

            elif pdf is None:
                st.markdown(f'<h5 style="text-align: center;color: orange;">Please Upload Your Resume</h5>', unsafe_allow_html=True)
            
            elif gemini_api_key == '':
                st.markdown(f'<h5 style="text-align: center;color: orange;">Gemini API Key not found in environment variables</h5>', unsafe_allow_html=True)


    def interview_questions_prompt(query_with_chunks, job_title):
        query = f'''Generate a comprehensive list of potential interview questions that might be asked during an interview for the position of {job_title}, based on the following resume.
        
        Structure your response in the following format:
        1. **Technical Questions**: Questions related to technical skills, tools, and technologies mentioned in the resume.
        2. **Experience-Based Questions**: Questions about past experiences, projects, and achievements.
        3. **Behavioral Questions**: Questions about soft skills, teamwork, problem-solving, and work style.
        4. **Role-Specific Questions**: Questions specific to the {job_title} position.
        5. **Company/Industry Questions**: Questions about industry knowledge and company fit.
        6. **Challenging Questions**: Difficult or unexpected questions that might be asked.
        
        For each question, provide a brief explanation of why it might be asked and what the interviewer is looking to assess.
        Format each section with proper headings and bullet points for better readability.
        
        Resume Content:
        """""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""
        {query_with_chunks}
        """""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""
        '''
        return query


    def resume_interview_questions():

        with st.form(key='Interview Questions'):

            # User Upload the Resume
            add_vertical_space(1)
            pdf = st.file_uploader(label='Upload Your Resume', type='pdf')
            add_vertical_space(1)
            
            # User Input for Job Title
            job_title = st.text_input(label='Enter the Job Title you are interviewing for')
            add_vertical_space(1)

            # Hidden Gemini API Key
            gemini_api_key = GEMINI_API_KEY
            add_vertical_space(2)

            # Click on Submit Button
            submit = st.form_submit_button(label='Submit')
            add_vertical_space(1)
        
        add_vertical_space(3)
        if submit:
            if pdf is not None and job_title != '' and gemini_api_key != '':
                try:
                    with st.spinner('Generating interview questions...'):
                        pdf_chunks = resume_analyzer.pdf_to_chunks(pdf)
                        interview_prompt = resume_analyzer.interview_questions_prompt(query_with_chunks=pdf_chunks, job_title=job_title)
                        interview_questions = resume_analyzer.gemini(gemini_api_key=gemini_api_key, chunks=pdf_chunks, analyze=interview_prompt)

                    st.markdown(f'<h3 style="color: #673AB7; border-bottom: 2px solid #673AB7; padding-bottom: 8px;">Potential Interview Questions for {job_title}</h3>', unsafe_allow_html=True)
                    st.markdown(interview_questions, unsafe_allow_html=True)

                    # Create a download button for the interview questions
                    timestamp = time.strftime("%Y%m%d_%H%M%S")
                    filename = f"Interview_Questions_{job_title.replace(' ', '_')}_{timestamp}.txt"
                    
                    # Format the content for the file
                    file_content = f"# Potential Interview Questions for {job_title}\n\n"
                    file_content += f"Generated on: {time.strftime('%Y-%m-%d %H:%M:%S')}\n\n"
                    file_content += interview_questions
                    
                    # Create a download button
                    st.download_button(
                        label="Download Interview Questions",
                        data=file_content,
                        file_name=filename,
                        mime="text/plain"
                    )

                except Exception as e:
                    st.markdown(f'<h5 style="text-align: center;color: orange;">{e}</h5>', unsafe_allow_html=True)

            elif pdf is None:
                st.markdown(f'<h5 style="text-align: center;color: orange;">Please Upload Your Resume</h5>', unsafe_allow_html=True)
            
            elif job_title == '':
                st.markdown(f'<h5 style="text-align: center;color: orange;">Please Enter the Job Title</h5>', unsafe_allow_html=True)
            
            elif gemini_api_key == '':
                st.markdown(f'<h5 style="text-align: center;color: orange;">Gemini API Key not found in environment variables</h5>', unsafe_allow_html=True)


    def cv_analysis_prompt(query_with_chunks):
        query = f'''Provide a comprehensive analysis of the following CV (Curriculum Vitae), focusing on academic and research backgrounds.
        
        Structure your response in the following format:
        1. **Academic Profile**: Analyze the educational background, degrees, institutions, and academic achievements.
        2. **Research Experience**: Evaluate research projects, methodologies, and contributions to the field.
        3. **Publications & Presentations**: Assess any published works, conference presentations, or academic contributions.
        4. **Teaching & Mentorship**: Analyze teaching experience, course development, and student mentorship.
        5. **Technical & Specialized Skills**: Identify specialized technical skills, laboratory techniques, or domain expertise.
        6. **Academic Achievements & Awards**: Highlight scholarships, grants, awards, and other academic recognitions.
        7. **Professional Development**: Evaluate professional memberships, certifications, and continuing education.
        8. **Recommendations for Academic/Research Positions**: Provide specific suggestions to strengthen the CV for academic or research roles.
        
        Format each section with proper headings and bullet points for better readability.
        Provide specific, actionable feedback that can help improve the CV for academic or research positions.
        
        CV Content:
        """""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""
        {query_with_chunks}
        """""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""
        '''
        return query

    def cv_analysis():
        with st.form(key='CV Analysis'):
            # User Upload the CV
            add_vertical_space(1)
            pdf = st.file_uploader(label='Upload Your CV', type='pdf')
            add_vertical_space(1)

            # Hidden Gemini API Key
            gemini_api_key = GEMINI_API_KEY
            add_vertical_space(2)

            # Click on Submit Button
            submit = st.form_submit_button(label='Submit')
            add_vertical_space(1)
        
        add_vertical_space(3)
        if submit:
            if pdf is not None and gemini_api_key != '':
                try:
                    with st.spinner('Processing...'):
                        pdf_chunks = resume_analyzer.pdf_to_chunks(pdf)
                        cv_prompt = resume_analyzer.cv_analysis_prompt(query_with_chunks=pdf_chunks)
                        cv_analysis = resume_analyzer.gemini(gemini_api_key=gemini_api_key, chunks=pdf_chunks, analyze=cv_prompt)

                    st.markdown(f'<h3 style="color: #3F51B5; border-bottom: 2px solid #3F51B5; padding-bottom: 8px;">CV Analysis for Academic/Research Positions</h3>', unsafe_allow_html=True)
                    st.markdown(cv_analysis, unsafe_allow_html=True)

                except Exception as e:
                    st.markdown(f'<h5 style="text-align: center;color: orange;">{e}</h5>', unsafe_allow_html=True)

            elif pdf is None:
                st.markdown(f'<h5 style="text-align: center;color: orange;">Please Upload Your CV</h5>', unsafe_allow_html=True)
            
            elif gemini_api_key == '':
                st.markdown(f'<h5 style="text-align: center;color: orange;">Gemini API Key not found in environment variables</h5>', unsafe_allow_html=True)

    def create_resume_prompt(personal_info, education, work_experience, skills, achievements, resume_type):
        prompt = f"""
        Create a professional resume based on the following information:
        
        # Personal Information
        {personal_info}
        
        # Education
        {education}
        
        # Work Experience
        {work_experience}
        
        # Skills
        {skills}
        
        # Achievements & Certifications
        {achievements}
        
        Resume Type: {resume_type}
        
        Format the resume in a clean, professional way using markdown formatting. Include appropriate sections and organize the information logically. Make it ATS-friendly if that's the selected resume type.
        """
        return prompt
    
    def create_resume():
        st.markdown(f'<h3 style="color: #2196F3; border-bottom: 2px solid #2196F3; padding-bottom: 8px;">Resume Creator</h3>', unsafe_allow_html=True)
        
        with st.expander("How to use this tool", expanded=True):
            st.markdown("""
            This tool helps you create a professional resume from scratch. Follow these steps:
            1. Fill in your details in each section below
            2. Select the type of resume you want to create
            3. Click 'Generate Resume' to create your resume
            4. Download the generated resume as a PDF, Word document, or text file
            
            **Tips for best results:**
            - Be specific about your achievements and use numbers when possible
            - Include relevant skills for your target job
            - Provide detailed work experience with dates, job titles, and responsibilities
            """)
        
        # Create PDF version of the resume
        def create_pdf(resume_text, name):
            from reportlab.lib.pagesizes import letter
            from reportlab.lib import colors
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.enums import TA_CENTER
            from reportlab.lib.units import inch
            from io import BytesIO
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter, 
                                   rightMargin=72, leftMargin=72,
                                   topMargin=72, bottomMargin=72)
            
            styles = getSampleStyleSheet()
            styles.add(ParagraphStyle(name='Heading1', 
                                     parent=styles['Heading1'], 
                                     fontSize=16, 
                                     alignment=TA_CENTER,
                                     spaceAfter=12))
            styles.add(ParagraphStyle(name='Heading2', 
                                     parent=styles['Heading2'], 
                                     fontSize=14, 
                                     spaceAfter=10))
            styles.add(ParagraphStyle(name='Normal', 
                                     parent=styles['Normal'], 
                                     fontSize=11, 
                                     spaceAfter=8))
            
            # Process the resume text to create a structured document
            elements = []
            
            # Split the resume text into lines
            lines = resume_text.split('\n')
            
            # Process each line
            current_section = None
            section_content = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                # Check if line is a heading (starts with # or ##)
                if line.startswith('# '):
                    # Add the previous section if it exists
                    if current_section and section_content:
                        elements.append(Paragraph(current_section, styles['Heading1']))
                        for content in section_content:
                            elements.append(Paragraph(content, styles['Normal']))
                        elements.append(Spacer(1, 12))
                    
                    # Start a new section
                    current_section = line[2:].strip()
                    section_content = []
                elif line.startswith('## '):
                    # Add the previous section if it exists
                    if current_section and section_content:
                        elements.append(Paragraph(current_section, styles['Heading1']))
                        for content in section_content:
                            elements.append(Paragraph(content, styles['Normal']))
                        elements.append(Spacer(1, 12))
                    
                    # Start a new section
                    current_section = line[3:].strip()
                    section_content = []
                elif line.startswith('**') and ':**' in line:
                    # This is likely a field name and value
                    field_parts = line.split(':**', 1)
                    if len(field_parts) == 2:
                        field_name = field_parts[0].replace('**', '')
                        field_value = field_parts[1].replace('**', '')
                        section_content.append(f"<b>{field_name}:</b> {field_value}")
                else:
                    # Regular content
                    section_content.append(line)
            
            # Add the last section
            if current_section and section_content:
                elements.append(Paragraph(current_section, styles['Heading1']))
                for content in section_content:
                    elements.append(Paragraph(content, styles['Normal']))
            
            # Build the PDF
            doc.build(elements)
            
            # Get the PDF data
            pdf_data = buffer.getvalue()
            buffer.close()
            
            return pdf_data
        
        # Create Word document version of the resume
        def create_word_doc(resume_text, name):
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches
            from io import BytesIO
            
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Split the resume text into lines
            lines = resume_text.split('\n')
            
            # Process each line
            current_section = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                # Check if line is a heading (starts with # or ##)
                if line.startswith('# '):
                    # Add as main heading
                    heading = doc.add_heading(line[2:].strip(), level=1)
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif line.startswith('## '):
                    # Add as subheading
                    heading = doc.add_heading(line[3:].strip(), level=2)
                elif line.startswith('**') and ':**' in line:
                    # This is likely a field name and value
                    field_parts = line.split(':**', 1)
                    if len(field_parts) == 2:
                        field_name = field_parts[0].replace('**', '')
                        field_value = field_parts[1].replace('**', '')
                        p = doc.add_paragraph()
                        p.add_run(f"{field_name}: ").bold = True
                        p.add_run(field_value)
                else:
                    # Regular content
                    doc.add_paragraph(line)
            
            # Save to a BytesIO object
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Get the Word document data
            docx_data = buffer.getvalue()
            buffer.close()
            
            return docx_data
        
        # NEW SIMPLIFIED RESUME CREATION APPROACH
        
        # Resume templates
        resume_templates = {
            "Software Engineer": {
                "summary": "Experienced software engineer with expertise in developing scalable applications and solving complex problems. Skilled in multiple programming languages and frameworks with a focus on delivering high-quality code.",
                "education": "Bachelor of Science in Computer Science, University Name, Year\nRelevant coursework: Data Structures, Algorithms, Software Engineering",
                "experience": "Software Engineer, Company Name, Start Date - End Date\n- Developed and maintained web applications using [languages/frameworks]\n- Collaborated with cross-functional teams to deliver features on time\n- Improved application performance by X%\n\nJunior Developer, Previous Company, Start Date - End Date\n- Assisted in developing and testing new features\n- Participated in code reviews and agile development processes",
                "skills": "Programming Languages: Python, JavaScript, Java\nFrameworks: React, Django, Spring\nTools: Git, Docker, Jenkins\nSoft Skills: Problem-solving, Communication, Teamwork",
                "achievements": "- Received Employee of the Month award for exceptional performance\n- Contributed to open-source projects\n- Completed relevant certifications"
            },
            "Data Scientist": {
                "summary": "Results-driven data scientist with experience in statistical analysis, machine learning, and data visualization. Skilled in extracting insights from complex datasets and communicating findings to stakeholders.",
                "education": "Master of Science in Data Science, University Name, Year\nBachelor of Science in Statistics, University Name, Year",
                "experience": "Data Scientist, Company Name, Start Date - End Date\n- Developed machine learning models to predict customer behavior\n- Created data visualizations to communicate insights to stakeholders\n- Improved prediction accuracy by X%\n\nData Analyst, Previous Company, Start Date - End Date\n- Performed statistical analysis on large datasets\n- Generated reports and dashboards for business users",
                "skills": "Programming: Python, R, SQL\nTools: Pandas, NumPy, Scikit-learn, TensorFlow\nVisualization: Tableau, Power BI, Matplotlib\nSoft Skills: Critical thinking, Communication, Problem-solving",
                "achievements": "- Published research paper on [topic]\n- Completed relevant certifications\n- Won hackathon for innovative data solution"
            },
            "Marketing Professional": {
                "summary": "Creative marketing professional with experience in digital marketing, campaign management, and brand development. Proven track record of increasing engagement and driving conversions through strategic marketing initiatives.",
                "education": "Bachelor of Business Administration in Marketing, University Name, Year",
                "experience": "Marketing Manager, Company Name, Start Date - End Date\n- Developed and executed marketing campaigns across multiple channels\n- Managed a team of X marketing specialists\n- Increased website traffic by X% and conversions by Y%\n\nMarketing Specialist, Previous Company, Start Date - End Date\n- Assisted in creating content for social media and email campaigns\n- Analyzed campaign performance and provided recommendations",
                "skills": "Digital Marketing: SEO, SEM, Social Media Marketing\nTools: Google Analytics, HubSpot, Mailchimp\nContent Creation: Copywriting, Graphic Design basics\nSoft Skills: Creativity, Communication, Project Management",
                "achievements": "- Received award for most successful marketing campaign\n- Increased social media following by X%\n- Completed relevant certifications"
            },
            "Project Manager": {
                "summary": "Detail-oriented project manager with experience in leading cross-functional teams and delivering projects on time and within budget. Skilled in project planning, risk management, and stakeholder communication.",
                "education": "Bachelor of Business Administration, University Name, Year\nProject Management Professional (PMP) Certification",
                "experience": "Project Manager, Company Name, Start Date - End Date\n- Led cross-functional teams to deliver X projects on time and within budget\n- Managed project budgets totaling $X\n- Implemented project management methodologies to improve efficiency\n\nAssistant Project Manager, Previous Company, Start Date - End Date\n- Assisted in project planning and execution\n- Coordinated with team members and stakeholders",
                "skills": "Project Management: Agile, Scrum, Waterfall\nTools: JIRA, Microsoft Project, Asana\nBudgeting and Resource Allocation\nSoft Skills: Leadership, Communication, Problem-solving",
                "achievements": "- Completed PMP certification\n- Successfully delivered high-impact project that increased revenue by X%\n- Received recognition for team leadership"
            },
            "Custom": {
                "summary": "",
                "education": "",
                "experience": "",
                "skills": "",
                "achievements": ""
            }
        }
        
        # Template selection
        template_options = list(resume_templates.keys())
        selected_template = st.selectbox("Choose a resume template or start from scratch", template_options)
        
        # Get the selected template
        template = resume_templates[selected_template]
        
        # Create tabs for different sections
        tab1, tab2, tab3, tab4, tab5 = st.tabs(["Personal Info", "Experience", "Education", "Skills", "Generate"])
        
        with tab1:
            st.subheader("Personal Information")
            
            # Two columns for name and contact
            col1, col2 = st.columns(2)
            with col1:
                name = st.text_input("Full Name")
            with col2:
                contact = st.text_input("Email")
                
            phone = st.text_input("Phone Number")
            linkedin = st.text_input("LinkedIn URL", help="A LinkedIn profile URL is highly recommended for professional networking")
            github = st.text_input("GitHub URL (optional)", help="A GitHub profile is recommended for technical roles to showcase your projects")
            
            st.subheader("Professional Summary")
            summary = st.text_area("Professional Summary", value=template["summary"], height=150)
            
            # Add resume tips in an expander
            with st.expander("Resume Writing Tips", expanded=False):
                st.markdown("""
                ### Key Resume Writing Tips:
                
                * **Contact Information:** Include all essential contact details. LinkedIn and/or GitHub profiles are highly recommended.
                * **Quantifiable Results:** Wherever possible, use numbers and metrics to demonstrate the impact of your work (e.g., percentage increases, cost savings, etc.).
                * **Specificity in Experience:** Include specific technologies used and project details to add context and make your experience more compelling.
                * **Action Verbs:** Start each bullet point with a strong action verb (e.g., developed, managed, increased, analyzed).
                * **Keywords:** Tailor your resume to the specific jobs you're applying for by including relevant keywords from the job descriptions.
                * **Consistency:** Use consistent formatting throughout the resume (e.g., bolding, italics, capitalization).
                * **Clean and Concise Language:** Use concise language and avoid jargon where possible.
                * **Length:** Aim to keep your resume to one page, especially if you have less than 10 years of experience.
                * **File Format:** Save your resume as a PDF to preserve formatting.
                
                Remember to customize your resume further for each specific job application!
                """)
        
        with tab2:
            st.subheader("Work Experience")
            st.markdown("List your work experience with most recent first")
            
            experience = st.text_area("Work Experience", value=template["experience"], height=300, 
                                    help="Format: Job Title, Company, Dates\n- Responsibility/Achievement\n- Responsibility/Achievement\n\nTIP: Start each bullet with an action verb and include quantifiable results where possible.")
            
            # Add action verbs suggestion box
            with st.expander("Suggested Action Verbs", expanded=False):
                st.markdown("""
                ### Effective Action Verbs for Resume Bullet Points:
                
                **Leadership:** Led, Managed, Directed, Coordinated, Oversaw, Supervised, Guided, Spearheaded
                
                **Achievement:** Achieved, Improved, Increased, Reduced, Decreased, Saved, Generated, Delivered
                
                **Development:** Developed, Created, Designed, Established, Implemented, Built, Launched, Initiated
                
                **Analysis:** Analyzed, Evaluated, Assessed, Researched, Identified, Solved, Investigated, Examined
                
                **Communication:** Presented, Communicated, Negotiated, Collaborated, Consulted, Facilitated, Persuaded
                
                **Technical:** Programmed, Engineered, Coded, Debugged, Tested, Maintained, Upgraded, Optimized
                
                Remember to be specific about your achievements and use numbers when possible!
                """)
        
        with tab3:
            st.subheader("Education")
            education = st.text_area("Education Details", value=template["education"], height=200,
                                   help="Format: Degree, Institution, Year\nRelevant coursework or achievements")
            
        with tab4:
            st.subheader("Skills & Achievements")
            
            skills = st.text_area("Skills", value=template["skills"], height=150,
                                help="List technical skills, soft skills, and tools you're proficient with. Include keywords relevant to your target job.")
            
            achievements = st.text_area("Achievements & Certifications", value=template["achievements"], height=150,
                                      help="List notable achievements, certifications, or awards. Quantify results where possible.")
            
            # Resume style options
            resume_style = st.radio("Resume Style", 
                                  ["Professional (Traditional)", "Modern", "ATS-Optimized"],
                                  horizontal=True)
            
            # Add ATS tips
            with st.expander("ATS Optimization Tips", expanded=False):
                st.markdown("""
                ### Tips for ATS-Friendly Resumes:
                
                * **Use Standard Section Headers:** "Experience," "Education," "Skills," etc.
                * **Include Keywords:** Mirror language from the job description
                * **Avoid Tables and Graphics:** These can confuse ATS systems
                * **Use Standard Fonts:** Stick with Arial, Calibri, or Times New Roman
                * **Simple Formatting:** Avoid text boxes, headers/footers, and complex layouts
                * **File Format:** Submit as a .docx or .pdf file
                * **Spell Check:** ATS systems may reject resumes with spelling errors
                
                The ATS-Optimized style option will format your resume to maximize compatibility with Applicant Tracking Systems.
                """)
        
        with tab5:
            st.subheader("Generate Your Resume")
            st.markdown("Review your information and generate your resume")
            
            # Hidden Gemini API Key
            gemini_api_key = GEMINI_API_KEY
            
            # Generate button
            generate_button = st.button("Generate Resume", type="primary", use_container_width=True)
            
            if generate_button:
                if gemini_api_key != '':
                    # Prepare the information
                    personal_info = f"Name: {name}\nEmail: {contact}\nPhone: {phone}"
                    if linkedin:
                        personal_info += f"\nLinkedIn: {linkedin}"
                    if github:
                        personal_info += f"\nGitHub: {github}"
                    personal_info += f"\nSummary: {summary}"
                    
                    try:
                        with st.spinner('Creating your professional resume...'):
                            # Create the prompt with selected style
                            style_instruction = ""
                            if resume_style == "Professional (Traditional)":
                                style_instruction = "Create a traditional, professional resume format."
                            elif resume_style == "Modern":
                                style_instruction = "Create a modern, visually appealing resume format."
                            elif resume_style == "ATS-Optimized":
                                style_instruction = "Create an ATS-optimized resume that will pass through applicant tracking systems."
                            
                            create_resume_prompt_text = f"""
                            Create a professional resume based on the following information:
                            
                            # Personal Information
                            {personal_info}
                            
                            # Education
                            {education}
                            
                            # Work Experience
                            {experience}
                            
                            # Skills
                            {skills}
                            
                            # Achievements & Certifications
                            {achievements}
                            
                            {style_instruction}
                            
                            Format the resume in a clean, professional way using markdown formatting. Include appropriate sections and organize the information logically.
                            """
                            
                            # Use a simplified approach since we don't need to process PDF chunks
                            generation_config = {
                                "temperature": 0.7,
                                "top_p": 1,
                                "top_k": 1,
                                "max_output_tokens": 2000,
                            }
                            
                            model = genai.GenerativeModel(
                                model_name="models/gemini-1.5-pro-latest",
                                generation_config=generation_config,
                            )
                            
                            # Improved response handling with error checking
                            try:
                                response = model.generate_content(create_resume_prompt_text)
                                generated_resume = response.text
                            except Exception as e:
                                # Fallback content if the model response is invalid
                                generated_resume = f"""
# Professional Resume

## Personal Information
**Name:** {name}
**Email:** {contact}
**Phone:** {phone}
{"**LinkedIn:** " + linkedin if linkedin else ""}
{"**GitHub:** " + github if github else ""}

## Professional Summary
{summary}

## Education
{education}

## Work Experience
{experience}

## Skills
{skills}

## Achievements & Certifications
{achievements}
"""
                                st.warning(f"The AI model encountered an issue: {str(e)}. A basic template has been created instead.")
                        
                            # Display the generated resume
                            st.markdown(f'<h3 style="color: #4CAF50; border-bottom: 2px solid #4CAF50; padding-bottom: 8px;">Your Generated Resume</h3>', unsafe_allow_html=True)
                            
                            # Create a container with a border for the resume
                            st.markdown("""
                            <style>
                            .resume-container {
                                border: 1px solid #ddd;
                                border-radius: 5px;
                                padding: 20px;
                                background-color: #f9f9f9;
                                margin-bottom: 20px;
                            }
                            </style>
                            """, unsafe_allow_html=True)
                            
                            st.markdown(f'<div class="resume-container">{generated_resume}</div>', unsafe_allow_html=True)
                            
                            # Generate the file name base
                            file_name_base = f"{name.replace(' ', '_')}_Resume" if name else "Generated_Resume"
                            
                            # Create columns for download buttons
                            col1, col2, col3 = st.columns(3)
                            
                            # Add download buttons for different formats
                            with col1:
                                st.download_button(
                                    label="Download as PDF",
                                    data=create_pdf(generated_resume, name),
                                    file_name=f"{file_name_base}.pdf",
                                    mime="application/pdf"
                                )
                            
                            with col2:
                                st.download_button(
                                    label="Download as Word",
                                    data=create_word_doc(generated_resume, name),
                                    file_name=f"{file_name_base}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                            
                            with col3:
                                st.download_button(
                                    label="Download as Text",
                                    data=generated_resume,
                                    file_name=f"{file_name_base}.txt",
                                    mime="text/plain"
                                )
                            
                            # Add copy button functionality with JavaScript
                            st.markdown("""
                            <div style="margin-top: 10px; margin-bottom: 20px;">
                                <button onclick="navigator.clipboard.writeText(document.querySelector('.resume-container').innerText);"
                                        style="background-color: #2196F3; color: white; border: none; padding: 10px 15px; 
                                        border-radius: 4px; cursor: pointer; width: 100%;">
                                    Copy Resume to Clipboard
                                </button>
                            </div>
                            """, unsafe_allow_html=True)
                            
                            # Tips for using the resume
                            with st.expander("Tips for using your new resume"):
                                st.markdown("""
                                ### Next Steps:
                                1. **Review and Edit**: Check for any inaccuracies or areas that need customization
                                2. **Formatting**: The PDF and Word versions maintain better formatting than plain text
                                3. **Customization**: Tailor your resume for specific job applications
                                4. **ATS Check**: Consider running your resume through an ATS checker before applying
                                5. **Professional Review**: Have someone in your industry review your resume for feedback
                                
                                ### Key Resume Improvement Tips:
                                * **Contact Information:** Include essential contact details with LinkedIn and/or GitHub profiles
                                * **Quantifiable Results:** Use numbers and metrics to demonstrate impact (e.g., "Improved performance by X%")
                                * **Specificity in Experience:** Include specific technologies used and project details
                                * **Action Verbs:** Start bullet points with strong action verbs
                                * **Keywords:** Include relevant keywords from job descriptions
                                * **Consistency:** Use consistent formatting for dates, bullet points, and section headings
                                * **Clean and Concise Language:** Use concise language and avoid jargon where possible
                                * **Length:** Aim for one page if you have less than 10 years of experience
                                * **File Format:** Use PDF format to preserve formatting
                                
                                Remember to replace any placeholder text with your specific information and tailor the resume to match the requirements of the specific job you're applying for.
                                """)
                    
                    except Exception as e:
                        st.markdown(f'<h5 style="text-align: center;color: orange;">{e}</h5>', unsafe_allow_html=True)
                
                elif gemini_api_key == '':
                    st.markdown(f'<h5 style="text-align: center;color: orange;">Gemini API Key not found in environment variables</h5>', unsafe_allow_html=True)


class linkedin_scraper:

    def webdriver_setup():
            
        options = webdriver.ChromeOptions()
        options.add_argument('--headless')
        options.add_argument('--no-sandbox')
        options.add_argument('--disable-dev-shm-usage')
        # Add performance improvements
        options.add_argument('--disable-gpu')
        options.add_argument('--disable-extensions')
        options.add_argument('--disable-infobars')
        options.add_argument('--disable-notifications')
        options.add_argument('--blink-settings=imagesEnabled=false')  # Disable images
        options.add_argument('--disable-javascript')  # Disable JavaScript where possible
        options.add_argument('--incognito')  # Use incognito mode to avoid caching issues

        driver = webdriver.Chrome(options=options)
        driver.set_page_load_timeout(30)  # Set timeout to avoid hanging
        driver.maximize_window()
        return driver


    def get_userinput():

        add_vertical_space(2)
        with st.form(key='linkedin_scarp'):

            add_vertical_space(1)
            col1,col2,col3 = st.columns([0.5,0.3,0.2], gap='medium')
            with col1:
                job_title_input = st.text_input(label='Job Title')
                job_title_input = job_title_input.split(',')
            with col2:
                job_location = st.text_input(label='Job Location', value='India')
            with col3:
                job_count = st.number_input(label='Job Count', min_value=1, value=1, step=1)

            # Submit Button
            add_vertical_space(1)
            submit = st.form_submit_button(label='Submit')
            add_vertical_space(1)
        
        return job_title_input, job_location, job_count, submit


    def build_url(job_title, job_location):

        b = []
        for i in job_title:
            x = i.split()
            y = '%20'.join(x)
            b.append(y)

        job_title = '%2C%20'.join(b)
        link = f"https://in.linkedin.com/jobs/search?keywords={job_title}&location={job_location}&locationId=&geoId=102713980&f_TPR=r604800&position=1&pageNum=0"

        return link
    

    def open_link(driver, link):
        # Set a maximum number of retries to avoid infinite loops
        max_retries = 3
        retries = 0
        
        while retries < max_retries:
            try:
                driver.get(link)
                # Reduced wait time
                driver.implicitly_wait(3)
                time.sleep(1)  # Reduced sleep time
                driver.find_element(by=By.CSS_SELECTOR, value='span.switcher-tabs__placeholder-text.m-auto')
                return True
            
            except NoSuchElementException:
                retries += 1
                if retries >= max_retries:
                    st.warning(f"Failed to load page after {max_retries} attempts. Continuing with limited functionality.")
                    return False
                continue


    def link_open_scrolldown(driver, link, job_count):
        
        # Open the Link in LinkedIn
        success = linkedin_scraper.open_link(driver, link)
        if not success:
            return
            
        # Calculate scroll iterations based on job count
        # More jobs require more scrolling, but we can optimize
        scroll_iterations = min(job_count * 2, 10)  # Cap at 10 to avoid excessive scrolling
        
        # Scroll Down the Page more efficiently
        for i in range(scroll_iterations):
            # Dismiss sign-in modal if present
            try:
                driver.find_element(by=By.CSS_SELECTOR, 
                                value="button[data-tracking-control-name='public_jobs_contextual-sign-in-modal_modal_dismiss']>icon>svg").click()
            except:
                pass

            # Scroll down in larger increments
            driver.execute_script(f"window.scrollBy(0, {1000});")
            
            # Click on See More Jobs Button if Present, but don't wait too long
            try:
                driver.find_element(by=By.CSS_SELECTOR, value="button[aria-label='See more jobs']").click()
                driver.implicitly_wait(2)  # Reduced wait time
            except:
                pass
            
            # Short pause between scrolls
            time.sleep(0.5)  # Reduced sleep time


    def job_title_filter(scrap_job_title, user_job_title_input):
        
        # User Job Title Convert into Lower Case
        user_input = [i.lower().strip() for i in user_job_title_input]

        # scraped Job Title Convert into Lower Case
        scrap_title = scrap_job_title.lower().strip()

        # Verify Any User Job Title in the scraped Job Title
        for i in user_input:
            # Check if all words in the user input are in the scraped title
            if all(word in scrap_title for word in i.split()):
                return scrap_job_title
                
        return np.nan


    def scrap_company_data(driver, job_title_input, job_location):

        try:
            # Get all job cards at once
            job_cards = driver.find_elements(by=By.CSS_SELECTOR, value='.base-card')
            
            company_names = []
            job_titles = []
            locations = []
            urls = []
            
            # Process each job card
            for card in job_cards:
                try:
                    # Extract data from each card
                    company_name = card.find_element(by=By.CSS_SELECTOR, value='.base-search-card__subtitle').text
                    job_title = card.find_element(by=By.CSS_SELECTOR, value='.base-search-card__title').text
                    location = card.find_element(by=By.CSS_SELECTOR, value='.job-search-card__location').text
                    url = card.find_element(by=By.CSS_SELECTOR, value='a').get_attribute('href')
                    
                    company_names.append(company_name)
                    job_titles.append(job_title)
                    locations.append(location)
                    urls.append(url)
                except:
                    # Skip cards with missing data
                    continue
            
            # Create DataFrame
            df = pd.DataFrame({
                'Company Name': company_names,
                'Job Title': job_titles,
                'Location': locations,
                'Website URL': urls
            })
            
            # Apply filters more efficiently
            # Filter job titles
            df['Job Title'] = df['Job Title'].apply(lambda x: linkedin_scraper.job_title_filter(x, job_title_input))
            
            # Filter locations
            job_location_lower = job_location.lower()
            df['Location'] = df['Location'].apply(lambda x: x if job_location_lower in x.lower() else np.nan)
            
            # Drop rows with NaN values and reset index
            df = df.dropna()
            df.reset_index(drop=True, inplace=True)
            
            return df
            
        except Exception as e:
            st.error(f"Error scraping job data: {e}")
            return pd.DataFrame(columns=['Company Name', 'Job Title', 'Location', 'Website URL'])


    def scrap_job_description(driver, df, job_count):
        
        # Get URL into List
        website_url = df['Website URL'].tolist()
        
        # Limit to requested job count
        website_url = website_url[:job_count]
        
        # Create progress bar
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        # Scrap the Job Description
        job_description = []
        
        for i, url in enumerate(website_url):
            try:
                # Update progress
                progress = int((i + 1) / len(website_url) * 100)
                progress_bar.progress(progress)
                status_text.text(f"Processing job {i+1} of {len(website_url)}")
                
                # Open the Link in LinkedIn with timeout
                driver.set_page_load_timeout(15)  # Set shorter timeout for each job page
                driver.get(url)
                driver.implicitly_wait(3)  # Reduced wait time
                
                # Try to click Show More button with a short timeout
                try:
                    show_more = driver.find_element(by=By.CSS_SELECTOR, value='button[data-tracking-control-name="public_jobs_show-more-html-btn"]')
                    driver.execute_script("arguments[0].click();", show_more)  # Use JavaScript click which is more reliable
                except:
                    pass  # Continue even if button not found
                
                # Get Job Description with a short timeout
                try:
                    description = driver.find_element(by=By.CSS_SELECTOR, value='div[class="show-more-less-html__markup relative overflow-hidden"]')
                    data = description.text
                    
                    # Check Description length
                    if len(data.strip()) > 0:
                        job_description.append(data)
                    else:
                        job_description.append('Description Not Available')
                except:
                    job_description.append('Description Not Available')
            
            except Exception as e:
                job_description.append('Description Not Available')
                st.warning(f"Could not process job {i+1}: {str(e)[:100]}...")
        
        # Clear progress indicators
        progress_bar.empty()
        status_text.empty()
        
        # Filter the Job Description
        df = df.iloc[:len(job_description), :]
        
        # Add Job Description in Dataframe
        df['Job Description'] = pd.DataFrame(job_description, columns=['Description'])
        df['Job Description'] = df['Job Description'].apply(lambda x: np.nan if x=='Description Not Available' else x)
        df = df.dropna()
        df.reset_index(drop=True, inplace=True)
        return df


    def display_data_userinterface(df_final):
        # Display the Data in User Interface
        add_vertical_space(1)
        
        if len(df_final) > 0:
            # Add download button for CSV
            csv = df_final.to_csv(index=False)
            st.download_button(
                label="Download job data as CSV",
                data=csv,
                file_name="linkedin_jobs.csv",
                mime="text/csv",
            )
            
            # Use tabs for better organization
            job_tabs = st.tabs([f"Job {i+1}" for i in range(len(df_final))])
            
            for i, tab in enumerate(job_tabs):
                with tab:
                    col1, col2 = st.columns([2, 1])
                    with col1:
                        st.subheader(df_final.iloc[i,1])  # Job Title as header
                        st.write(f"**Company:** {df_final.iloc[i,0]}")
                        st.write(f"**Location:** {df_final.iloc[i,2]}")
                    with col2:
                        st.markdown(f"[Apply on LinkedIn]({df_final.iloc[i,3]})")
                    
                    st.markdown("### Job Description")
                    st.write(df_final.iloc[i, 4])
                    st.divider()
        
        else:
            st.warning("No matching jobs found. Try different search terms or location.")


    def main():
        
        # Initially set driver to None
        driver = None
        
        try:
            job_title_input, job_location, job_count, submit = linkedin_scraper.get_userinput()
            add_vertical_space(2)
            
            if submit:
                if job_title_input != [''] and job_location != '':
                    
                    with st.spinner('Setting up browser...'):
                        driver = linkedin_scraper.webdriver_setup()
                                       
                    with st.spinner('Loading job listings...'):
                        # build URL based on User Job Title Input
                        link = linkedin_scraper.build_url(job_title_input, job_location)
                        # Open the Link in LinkedIn and Scroll Down the Page
                        linkedin_scraper.link_open_scrolldown(driver, link, job_count)

                    with st.spinner('Gathering job information...'):
                        # Scraping the Company Name, Location, Job Title and URL Data
                        df = linkedin_scraper.scrap_company_data(driver, job_title_input, job_location)
                        
                        if len(df) > 0:
                            # Scraping the Job Description Data
                            df_final = linkedin_scraper.scrap_job_description(driver, df, job_count)
                            # Display the Data in User Interface
                            linkedin_scraper.display_data_userinterface(df_final)
                        else:
                            st.warning("No jobs found matching your criteria. Try different search terms.")
                
                # If User Click Submit Button and Job Title is Empty
                elif job_title_input == ['']:
                    st.warning("Please enter a job title")
                
                elif job_location == '':
                    st.warning("Please enter a job location")

        except Exception as e:
            add_vertical_space(2)
            st.error(f"An error occurred: {str(e)}")
        
        finally:
            if driver:
                driver.quit()



# Streamlit Configuration Setup
streamlit_config()
add_vertical_space(2)



with st.sidebar:

    add_vertical_space(4)

    option = option_menu(menu_title='', options=['Summary', 'Strength', 'Weakness', 'Job Titles', 'Interview Questions', 'CV Analysis', 'Create Resume', 'Linkedin Jobs'],
                         icons=['house-fill', 'database-fill', 'pass-fill', 'list-ul', 'question-square', 'file-earmark-person', 'file-earmark-plus', 'linkedin'])

if option == 'Summary':

    resume_analyzer.resume_summary()

elif option == 'Strength':

    resume_analyzer.resume_strength()

elif option == 'Weakness':

    resume_analyzer.resume_weakness()

elif option == 'Job Titles':

    resume_analyzer.resume_job_title()

elif option == 'Interview Questions':

    resume_analyzer.resume_interview_questions()

elif option == 'CV Analysis':
    
    resume_analyzer.cv_analysis()

elif option == 'Create Resume':
    
    resume_analyzer.create_resume()

elif option == 'Linkedin Jobs':
    
    linkedin_scraper.main()
